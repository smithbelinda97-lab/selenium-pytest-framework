from docx import Document
from docx.shared import Inches


document = Document()

document.add_heading(
    "Automation Execution Report",
    level=1
)


def add_screenshot_to_report(
    step_name,
    screenshot_path
):

    document.add_heading(step_name, level=2)

    document.add_picture(
        screenshot_path,
        width=Inches(5.5)
    )

    document.add_paragraph(
        f"Screenshot captured for: {step_name}"
    )


def save_report(report_name):

    document.save(report_name)

    print(f"Report saved: {report_name}")